import requests
import openai
from urllib.request import urlopen
from urllib.parse import unquote
import json
from email.message import EmailMessage
from datetime import date
import smtplib
from docx import Document

# from email.mime.text import MIMEText

def generate_response_gpt3_5(prompt):
    response = openai.ChatCompletion.create(
        model ="gpt-3.5-turbo",
        messages=[
        {"role": "system", "content":"You are a helpful assistant."},
        {"role":"user","content":prompt}
        ]
    )
    print(response)
    summary = response["choices"][0]["message"]["content"].strip()
    return summary

def send_email(sender = "zokeechen@outlook.com", recipient = ["zokeechen@163.com"]):
    # recipient = ["yadongchen@tencent.com","zokeechen@163.com"]
    # gnews = open('newsdata.txt', 'r')
    # message = gnews.readlines()
    # # message = gnews.readlines()
    message = ""
    gnews = open('newsdata.txt', 'r')
    for line in gnews:
        message += line
        message += "\n"
    # print(message)
    gnews.close()
    try:
        email = EmailMessage()
        email["From"] = sender
        email["To"] = recipient
        email["Subject"] = "【ChatGPT每日新闻收录" + str(date.today()) + "】 — TEG战略发展中心推送(newsdata.io)"
        email.set_content(message)

        smtp = smtplib.SMTP("smtp-mail.outlook.com", port=587)
        smtp.starttls()
        smtp.login(sender, "")
        smtp.sendmail(sender, recipient, email.as_string())
        smtp.quit()
        print("Sent!")
    except:
        pass

# def add_hyperlink(paragraph, url, text, color, underline):
#     """
#     A function that places a hyperlink within a paragraph object.
#
#     :param paragraph: The paragraph we are adding the hyperlink to.
#     :param url: A string containing the required url
#     :param text: The text displayed for the url
#     :return: The hyperlink object
#     """
#
#     # This gets access to the document.xml.rels file and gets a new relation id value
#     part = paragraph.part
#     r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
#
#     # Create the w:hyperlink tag and add needed values
#     hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
#     hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )
#
#     # Create a w:r element
#     new_run = docx.oxml.shared.OxmlElement('w:r')
#
#     # Create a new w:rPr element
#     rPr = docx.oxml.shared.OxmlElement('w:rPr')
#
#     # Add color if it is given
#     if not color is None:
#       c = docx.oxml.shared.OxmlElement('w:color')
#       c.set(docx.oxml.shared.qn('w:val'), color)
#       rPr.append(c)
#
#     # Remove underlining if it is requested
#     if not underline:
#       u = docx.oxml.shared.OxmlElement('w:u')
#       u.set(docx.oxml.shared.qn('w:val'), 'none')
#       rPr.append(u)
#
#     # Join all the xml elements together add add the required text to the w:r element
#     new_run.append(rPr)
#     new_run.text = text
#     hyperlink.append(new_run)
#
#     paragraph._p.append(hyperlink)
#
#     return hyperlink

def get_analysis(url,topic):
    response = urlopen(url)
    data = json.loads(response.read().decode("utf-8"))
    articles = data["results"]
    gnews.write(f"关于【{unquote(topic)}】话题: \n")
    a = document.add_paragraph(f"关于【{unquote(topic)}】话题: ")
    for i in range(len(articles)):
        try:
            gnews.write(str(i+1))
            gnews.write(". \n")
            title = articles[i]['title']
            description = articles[i]['description']
            content = articles[i]['content']
            url = articles[i]['link']
        except:
            gnews.write("文章无法获取\n")
            pass

        try:
            prompt = f'''
            请总结下面新闻并翻译成中文发给我，根据网页内容，总结出两段文字，第一段是【标题】，第二段是【关键内容摘要】以最多4个要点形式展示详细内容摘要：Link:{url},Title: {title}, Description: {description}, Content: {content}
            '''
            prompt = f'''
            Please analyze the following news and send me back in language of Chinese，make sure your response is in chinese 中文，according to only the contents given, summarize to two paragraphs，first paragraph is Title，second paragraph is bullet point summary(at most 4 points)：Link:{url},Title: {title}, Description: {description}, Content: {content}
            '''
            prompt = f'''
            请确保你输出的内容是中文
            Your output should use the following template:
            Summary
            Facts
            Numbers(keep this section only with numbers given)
            -Bulletpoint
            Your task is to summarize the text I give you in up to seven concise bulletpoints and start with a short, high-quality summary. Your response should be in Simplified Chinese. Use the following text:
            Link:{url},
            Title: {title},
            Description: {description},
            Content: {content}
            请确保你输出的内容是中文
            '''
            print(prompt)
            summary = generate_response_gpt3_5(prompt)
            print(summary)
            # gnews.write("ChatGPT3.5 回答如下： \n")
            gnews.write(str(summary))
            gnews.write("\n")
            gnews.write("新闻链接: " + url)
            gnews.write("\n")
            gnews.write("-----------------")
            gnews.write("\n")
            a = document.add_paragraph(f"{i+1}")
            paragraph = document.add_paragraph("")
            paragraph.add_run(title).bold = True
            paragraph = document.add_paragraph(url)
            paragraph = document.add_paragraph(summary)
            a = document.add_paragraph("")

        except:
            gnews.write("ChatGPT 回答无法获取\n")
            pass


#-------------------------------------------------------------------------------------------------------
#设置OpenAI API密钥,news api
openai.api_key = "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"
newsdata_apikey = "pub_19244bc3ace0a1ddb9736bc0ba54f7203149c"
#Topic 设置为感兴趣的主题，当前 ChatGPT
# topic = "GPT%20OR%20OpenAI"
# url = 'https://newsdata.io/api/1/news?q='+topic+'&language=en&country=us&apikey=' + newsdata_apikey;
topics = ["ChatGPT","GPT","OpenAI","Generative%20AI","Large%20Language%20Model","New%20Bing", "Google%20Bard"]
# url = 'https://newsdata.io/api/1/news?qInTitle='+topic+'&language=en&country=us&apikey=' + newsdata_apikey;
#Head
gnews = open("newsdata.txt", "w")
gnews.write("Dear all, \n")
gnews.write("该报是由战略发展中心 AI 自动推送，帮助大家同步全球范围内 ChatGPT 每日新闻收录。仅在团队内部分享：\n\n")
gnews.write('''日报收录关键词包括"ChatGPT","GPT","OpenAI","Generative AI","Large Language Model","New Bing", "Google Bard"\n''')
document = Document()
for topic in topics:
    url = 'https://newsdata.io/api/1/news?qInTitle='+topic+'&language=en&country=us&apikey=' + newsdata_apikey
    get_analysis(url,topic)

#End
gnews.write(f"今日新闻到此结束")
gnews.close()
document.save('example.docx')

# #send email
# send_email()
